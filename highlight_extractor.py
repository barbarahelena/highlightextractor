#!/usr/bin/env python3
"""
PDF Highlight Extractor
Extracts highlighted text from PDF files and saves to markdown, docx, or txt format.
"""

import argparse
import re
import sys
from pathlib import Path

try:
    import fitz  # PyMuPDF
except ImportError:
    print("Error: PyMuPDF is required. Install it with: pip install PyMuPDF")
    sys.exit(1)


def extract_highlights(pdf_path: str) -> list[dict]:
    """
    Extract highlighted text from a PDF file.
    Returns a list of dicts with 'page' and 'text' keys.
    """
    highlights = []
    doc = fitz.open(pdf_path)

    for page_num, page in enumerate(doc, start=1):
        # Get all text with position info for better extraction
        text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)

        annotations = page.annots()
        if annotations is None:
            continue

        for annot in annotations:
            # Highlight annotation type is 8
            if annot.type[0] == 8:
                # Get text using multiple methods and pick the best one
                text = extract_highlight_text(page, annot, text_dict)

                if text:
                    cleaned = clean_text(text)
                    if cleaned and is_valid_text(cleaned):
                        highlights.append({
                            "page": page_num,
                            "text": cleaned,
                        })

    doc.close()
    return highlights


def extract_highlight_text(page, annot, text_dict) -> str:
    """Extract text from highlight annotation using multiple methods."""

    # Method 1: Try to get text from annotation's QuadPoints (most accurate)
    text = extract_from_quadpoints(page, annot)
    if text and is_valid_text(clean_text(text)):
        return text

    # Method 2: Fall back to rect-based extraction
    rect = annot.rect
    text = extract_text_from_rect(text_dict, rect)
    if text and is_valid_text(clean_text(text)):
        return text

    # Method 3: Simple rect extraction
    text = page.get_text("text", clip=rect).strip()
    return text


def extract_from_quadpoints(page, annot) -> str:
    """Extract text using annotation's QuadPoints for precise extraction."""
    try:
        # Get the quad points (4 corners per highlighted region)
        quads = annot.vertices
        if not quads:
            return ""

        text_parts = []

        # Process quads in groups of 4 points
        for i in range(0, len(quads), 4):
            if i + 3 >= len(quads):
                break

            # Get bounding rect from quad points
            points = quads[i:i+4]
            x_coords = [p[0] for p in points]
            y_coords = [p[1] for p in points]

            rect = fitz.Rect(
                min(x_coords),
                min(y_coords),
                max(x_coords),
                max(y_coords)
            )

            # Extract text from this rect
            text = page.get_text("text", clip=rect)
            if text:
                text_parts.append(text.strip())

        return " ".join(text_parts)
    except Exception:
        return ""


def extract_text_from_rect(text_dict, rect) -> str:
    """Extract text from text_dict that falls within the given rect."""
    text_parts = []

    for block in text_dict.get("blocks", []):
        if block.get("type") != 0:  # Only text blocks
            continue

        for line in block.get("lines", []):
            line_rect = fitz.Rect(line["bbox"])

            # Check if line intersects with highlight rect
            if not rect.intersects(line_rect):
                continue

            line_text = []
            for span in line.get("spans", []):
                span_rect = fitz.Rect(span["bbox"])

                # Check if span is mostly within the highlight rect
                if is_mostly_within(span_rect, rect, threshold=0.5):
                    line_text.append(span["text"])

            if line_text:
                text_parts.append("".join(line_text))

    return " ".join(text_parts)


def is_mostly_within(inner: fitz.Rect, outer: fitz.Rect, threshold: float = 0.5) -> bool:
    """Check if inner rect is mostly within outer rect."""
    if inner.is_empty or outer.is_empty:
        return False

    intersection = inner & outer  # Intersection
    if intersection.is_empty:
        return False

    inner_area = inner.width * inner.height
    if inner_area == 0:
        return False

    intersection_area = intersection.width * intersection.height
    return (intersection_area / inner_area) >= threshold


def clean_text(text: str) -> str:
    """Clean extracted text by removing artifacts and normalizing whitespace."""
    if not text:
        return ""

    # Replace common PDF ligatures with their expanded forms
    ligatures = {
        "\ufb00": "ff",  # ff ligature
        "\ufb01": "fi",  # fi ligature
        "\ufb02": "fl",  # fl ligature
        "\ufb03": "ffi", # ffi ligature
        "\ufb04": "ffl", # ffl ligature
        "\ufb05": "st",  # st ligature (long s)
        "\ufb06": "st",  # st ligature
    }
    for lig, replacement in ligatures.items():
        text = text.replace(lig, replacement)

    # Remove control characters except newline
    text = "".join(c for c in text if not (ord(c) < 32 and c not in "\n\t "))

    # Normalize different types of whitespace
    text = text.replace("\t", " ")
    text = text.replace("\r\n", " ")
    text = text.replace("\r", " ")
    text = text.replace("\n", " ")

    # Collapse multiple spaces
    text = re.sub(r" +", " ", text)

    # Fix hyphenation at line breaks (word- continuation -> word continuation)
    text = re.sub(r"(\w)-\s+(\w)", r"\1\2", text)

    # Fix common punctuation spacing issues
    text = re.sub(r"\s+([,.:;!?)])", r"\1", text)  # Remove space before punctuation
    text = re.sub(r"([(])\s+", r"\1", text)  # Remove space after opening paren

    # Fix common abbreviations that might have spaces
    text = re.sub(r"\bi\s*\.\s*e\s*\.", "i.e.", text)
    text = re.sub(r"\be\s*\.\s*g\s*\.", "e.g.", text)

    # Remove isolated single characters that are likely artifacts from ligatures
    # Pattern: space + single letter + space (but preserve "a", "A", uppercase "I")
    # These often come from ligatures (fi, fl, ff) being cut by highlight boundaries
    # Note: lowercase "i" between words is almost always an artifact, not the pronoun "I"
    text = re.sub(r"(?<=\s)[bcdefghijklmnopqrstuvwxyz](?=\s)", "", text)
    text = re.sub(r"(?<=\s)[BCDEFGHJKLMNOPQRSTUVWXYZ](?=\s)", "", text)

    # Remove isolated ligature fragments (ff, fi, fl, etc. surrounded by spaces)
    text = re.sub(r"(?<=\s)(ff|fi|fl|ffi|ffl)(?=\s)", "", text)

    # Remove leading single characters or ligature fragments at start of text
    text = re.sub(r"^[bcdefghijklmnopqrstuvwxyz]\s+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^(ff|fi|fl|ffi|ffl)\s+", "", text)

    # Remove trailing single characters at end
    text = re.sub(r"\s+[bcdefghijklmnopqrstuvwxyz]$", "", text, flags=re.IGNORECASE)

    # Collapse any resulting multiple spaces again
    text = re.sub(r" +", " ", text)

    # Fix double periods
    text = re.sub(r"\.\.+", ".", text)

    return text.strip()


def is_valid_text(text: str) -> bool:
    """Check if the text is valid (not just garbage characters)."""
    if not text or len(text) < 3:
        return False

    # Must contain at least one word with 2+ letters
    if not re.search(r"[a-zA-Z]{2,}", text):
        return False

    # Calculate ratio of letters to total non-whitespace characters
    non_space = [c for c in text if not c.isspace()]
    if not non_space:
        return False

    letters = sum(1 for c in non_space if c.isalpha())
    ratio = letters / len(non_space)

    # At least 40% should be letters
    return ratio > 0.4


def save_as_markdown(highlights: list[dict], output_path: str, pdf_name: str):
    """Save highlights to a markdown file."""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(f"# Highlights from {pdf_name}\n\n")

        current_page = None
        for h in highlights:
            if h["page"] != current_page:
                current_page = h["page"]
                f.write(f"\n## Page {current_page}\n\n")

            f.write(f"> {h['text']}\n\n")


def save_as_txt(highlights: list[dict], output_path: str, pdf_name: str):
    """Save highlights to a plain text file."""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(f"Highlights from {pdf_name}\n")
        f.write("=" * 50 + "\n\n")

        current_page = None
        for h in highlights:
            if h["page"] != current_page:
                current_page = h["page"]
                f.write(f"\n--- Page {current_page} ---\n\n")

            f.write(f"* {h['text']}\n\n")


def save_as_docx(highlights: list[dict], output_path: str, pdf_name: str):
    """Save highlights to a Word document."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        print("Error: python-docx is required for .docx output. Install it with: pip install python-docx")
        sys.exit(1)

    doc = Document()

    # Add title
    title = doc.add_heading(f"Highlights from {pdf_name}", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    current_page = None
    for h in highlights:
        if h["page"] != current_page:
            current_page = h["page"]
            doc.add_heading(f"Page {current_page}", level=1)

        para = doc.add_paragraph(style="Quote")
        para.add_run(h["text"])

    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(
        description="Extract highlighted text from PDF files",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python highlight_extractor.py document.pdf -f md
  python highlight_extractor.py document.pdf -f txt -o my_highlights.txt
  python highlight_extractor.py document.pdf -f docx
        """
    )

    parser.add_argument("pdf", help="Path to the PDF file")
    parser.add_argument(
        "-f", "--format",
        choices=["md", "txt", "docx"],
        default="md",
        help="Output format: md (markdown), txt (plain text), or docx (Word document). Default: md"
    )
    parser.add_argument(
        "-o", "--output",
        help="Output file path. If not specified, uses the PDF name with appropriate extension"
    )

    args = parser.parse_args()

    # Validate input file
    pdf_path = Path(args.pdf)
    if not pdf_path.exists():
        print(f"Error: File not found: {args.pdf}")
        sys.exit(1)

    if not pdf_path.suffix.lower() == ".pdf":
        print("Warning: File does not have .pdf extension")

    # Determine output path
    if args.output:
        output_path = args.output
    else:
        ext_map = {"md": ".md", "txt": ".txt", "docx": ".docx"}
        output_path = pdf_path.stem + "_highlights" + ext_map[args.format]

    # Extract highlights
    print(f"Extracting highlights from: {pdf_path.name}")
    highlights = extract_highlights(str(pdf_path))

    if not highlights:
        print("No highlights found in the PDF.")
        sys.exit(0)

    print(f"Found {len(highlights)} highlight(s)")

    # Save to chosen format
    pdf_name = pdf_path.name
    if args.format == "md":
        save_as_markdown(highlights, output_path, pdf_name)
    elif args.format == "txt":
        save_as_txt(highlights, output_path, pdf_name)
    elif args.format == "docx":
        save_as_docx(highlights, output_path, pdf_name)

    print(f"Highlights saved to: {output_path}")


if __name__ == "__main__":
    main()
